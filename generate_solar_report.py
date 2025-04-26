import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
import os
import openai
from dotenv import load_dotenv
import json
import re

# --- Load OpenAI API key ---
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

# --- Read the analysis results ---
results_dir = "solar_results"
results_file = os.path.join(results_dir, 'solar_analysis_results.csv')
summary_file = os.path.join(results_dir, 'solar_summary.json')
if not os.path.exists(results_file):
    raise FileNotFoundError(f"{results_file} not found. Please run solar_calculator.py first.")
if not os.path.exists(summary_file):
    raise FileNotFoundError(f"{summary_file} not found. Please run solar_calculator.py first.")

df = pd.read_csv(results_file, parse_dates=["datetime"])
with open(summary_file) as f:
    summary_data = json.load(f)

# --- Prepare summary tables for both scenarios ---
solar_only = summary_data["solar_only"]
battery = summary_data["battery"]
supply_charge_total = summary_data["supply_charge_total"]
system_cost = summary_data["system_cost"]

energy_rows_solar = [
    ("Total energy used (kWh)", f"{solar_only['total_usage']:.2f}"),
    ("Total solar produced (kWh)", f"{solar_only['total_solar']:.2f}"),
    ("Self-consumed solar (kWh, solar only)", f"{solar_only['total_self_consumed']:.2f}"),
    ("Exported to grid (kWh)", f"{solar_only['total_exported']:.2f}"),
    ("Imported from grid (kWh)", f"{solar_only['total_imported']:.2f}"),
]
cost_rows_solar = [
    ("Annual supply charge ($)", f"{supply_charge_total:,.2f}"),
    ("Total earned from exported electricity ($)", f"{solar_only['export_earnings']:,.2f}"),
    ("Cost without solar ($)", f"{solar_only['cost_without_solar']:,.2f}"),
    ("Cost with solar ($)", f"{solar_only['cost_with_solar']:,.2f}"),
    ("Total savings per year ($)", f"{solar_only['total_savings']:,.2f}"),
]
energy_rows_battery = [
    ("Total energy used (kWh)", f"{battery['total_usage']:.2f}"),
    ("Total solar produced (kWh)", f"{battery['total_solar']:.2f}"),
    ("Self-consumed solar (kWh, with battery)", f"{battery['total_self_consumed']:.2f}"),
    ("Exported to grid (kWh)", f"{battery['total_exported']:.2f}"),
    ("Imported from grid (kWh)", f"{battery['total_imported']:.2f}"),
]
if battery["use_battery"]:
    energy_rows_battery.append(("Total battery charge (kWh)", f"{battery['total_battery_charge']:.2f}"))
    energy_rows_battery.append(("Total battery discharge (kWh)", f"{battery['total_battery_discharge']:.2f}"))
cost_rows_battery = [
    ("Annual supply charge ($)", f"{supply_charge_total:,.2f}"),
    ("Total earned from exported electricity ($)", f"{battery['export_earnings']:,.2f}"),
    ("Cost without solar ($)", f"{battery['cost_without_solar']:,.2f}"),
    ("Cost with solar + battery ($)", f"{battery['cost_with_solar']:,.2f}"),
    ("Total savings per year ($)", f"{battery['total_savings']:,.2f}"),
    ("Estimated payback period (years)", f"{battery['payback_years']:.1f}"),
]

# --- GPT-4.1-mini summary and assumptions ---
def gpt_summary_and_assumptions(solar_only, battery, supply_charge_total):
    # Prepare variables for the prompt
    annual_solar_generation_kWh = battery['total_solar']
    annual_household_consumption_kWh = battery['total_usage']
    self_consumed_solar_kWh = solar_only['total_self_consumed']
    exported_solar_kWh = solar_only['total_exported']
    grid_import_kWh = solar_only['total_imported']
    annual_supply_charge = supply_charge_total
    annual_export_earnings = solar_only['export_earnings']
    total_cost_without_solar = solar_only['cost_without_solar']
    total_cost_with_solar = solar_only['cost_with_solar']
    annual_savings_solar_only = solar_only['total_savings']
    self_consumed_with_battery_kWh = battery['total_self_consumed']
    exported_with_battery_kWh = battery['total_exported']
    grid_import_with_battery_kWh = battery['total_imported']
    battery_charge_kWh = battery['total_battery_charge']
    battery_discharge_kWh = battery['total_battery_discharge']
    annual_export_earnings_with_battery = battery['export_earnings']
    total_cost_with_battery = battery['cost_with_solar']
    annual_savings_with_battery = battery['total_savings']
    payback_period_years = battery['payback_years']

    paragraphs = f'''
Paragraph 1: Solar Production and Basic Usage
The residential solar installation generates a substantial amount of energy annually, producing {annual_solar_generation_kWh:.2f} kWh of solar power. In the solar-only scenario, of the total {annual_household_consumption_kWh:.2f} kWh energy consumed, only {self_consumed_solar_kWh:.2f} kWh is self-consumed directly from solar, with the remainder exported to the grid at {exported_solar_kWh:.2f} kWh. The household imports {grid_import_kWh:.2f} kWh from the grid.

Paragraph 2: Financial Impact of Solar-Only Setup
Financially, this setup results in an annual supply charge of ${annual_supply_charge:,.2f}, with earnings of ${annual_export_earnings:,.2f} from exported electricity. This reduces the total electricity cost from ${total_cost_without_solar:,.2f} (without solar) to ${total_cost_with_solar:,.2f}, yielding yearly savings of ${annual_savings_solar_only:,.2f}.

Paragraph 3: Battery Storage Benefits
Introducing a battery storage system significantly enhances self-consumption, increasing it to {self_consumed_with_battery_kWh:.2f} kWh. This limits solar exports to {exported_with_battery_kWh:.2f} kWh and drastically reduces imports to just {grid_import_with_battery_kWh:.2f} kWh. The battery charges with {battery_charge_kWh:.2f} kWh and discharges {battery_discharge_kWh:.2f} kWh to support home energy use.

Paragraph 4: Financial Impact with Battery
Economically, despite slightly decreased earnings from export (${annual_export_earnings_with_battery:,.2f}), the battery-enabled system cuts electricity costs substantially to ${total_cost_with_battery:,.2f} annually, resulting in much higher savings of ${annual_savings_with_battery:,.2f}. An estimated payback period for the solar+battery system stands at {payback_period_years:.1f} years.

Paragraph 5: Final Comparison and Recommendations
Overall, the solar-only installation already offers significant cost and energy benefits by reducing grid reliance and electricity expenditures. However, adding battery storage maximises self-consumption, minimises grid dependency, and greatly boosts savings, justifying the additional investment over time. Both options contribute toward cleaner energy usage and improved household energy independence. Users should weigh initial costs, potential changes in export tariffs, and energy consumption patterns to decide the optimal system configuration.
'''

    prompt = f"""
You are an energy analyst. Using the following summary paragraphs, return a JSON object with the following keys: 'paragraph_1', 'paragraph_2', 'paragraph_3', 'paragraph_4', 'paragraph_5', and 'assumptions'. Each key should contain the corresponding paragraph as a string, with no headings or extra text. Do not include any headings in your response. The JSON should look like: {{"paragraph_1": "...", "paragraph_2": "...", ..., "assumptions": "..."}}.

{paragraphs}

Also, in the 'assumptions' field, briefly comment on:
- The DEBS electricity rebate used in the calculation could change in the future.
- Battery-to-grid export rebates are not included in the calculation.
- If the owner has an electric car, this would impact cost savings.
Encourage the reader to consider other scenarios that may affect the results.
"""
    response = openai.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=700
    )
    text = response.choices[0].message.content.strip()
    # Remove code block formatting if present
    if text.startswith('```'):
        text = text.strip('`\n')
        if text.startswith('json'):
            text = text[4:].strip()
    # Try to parse as JSON
    try:
        parsed = json.loads(text)
        return parsed, parsed.get("assumptions", "")
    except Exception:
        print("[Warning] GPT response was not valid JSON. Using fallback.")
        return {"paragraph_1": text}, ""

summary_json, assumptions_text = gpt_summary_and_assumptions(solar_only, battery, supply_charge_total)

# --- Create plots directory ---
plots_dir = "solar_report_plots"
os.makedirs(plots_dir, exist_ok=True)

# --- Plot 1: Hourly flows for a sample week ---
sample = df.iloc[:168]  # First week
plt.figure(figsize=(10, 6))
plt.plot(sample["datetime"], sample["usage_kwh"], label="Usage (kWh)")
plt.plot(sample["datetime"], sample["solar_kwh"], label="Solar (kWh)")
plt.plot(sample["datetime"], sample["self_consumed"], label="Self-consumed (kWh)")
plt.plot(sample["datetime"], sample["imported"], label="Grid Import (kWh)")
plt.plot(sample["datetime"], sample["exported"], label="Grid Export (kWh)")
if battery["use_battery"]:
    plt.plot(sample["datetime"], sample["battery_soc"], label="Battery SOC (kWh)")
plt.legend()
plt.title("Hourly Energy Flows (First Week)")
plt.xlabel("Datetime")
plt.ylabel("kWh")
plt.tight_layout()
plot1_path = os.path.join(plots_dir, "hourly_flows.png")
plt.savefig(plot1_path)
plt.close()

# --- Plot 2: Monthly totals (include battery charge/discharge if present) ---
df["month"] = df["datetime"].dt.month
gb = df.groupby("month").sum(numeric_only=True)
plt.figure(figsize=(10, 6))
plt.bar(gb.index-0.3, gb["usage_kwh"], width=0.15, label="Usage")
plt.bar(gb.index-0.15, gb["solar_kwh"], width=0.15, label="Solar")
plt.bar(gb.index, gb["self_consumed"], width=0.15, label="Self-consumed")
if battery["use_battery"]:
    plt.bar(gb.index+0.15, gb["battery_charge"], width=0.15, label="Battery Charge")
    plt.bar(gb.index+0.3, gb["battery_discharge"], width=0.15, label="Battery Discharge")
plt.title("Monthly Energy Totals")
plt.xlabel("Month")
plt.ylabel("kWh")
plt.legend()
plt.tight_layout()
plot2_path = os.path.join(plots_dir, "monthly_totals.png")
plt.savefig(plot2_path)
plt.close()

# --- Plot 3: Pie chart of energy breakdown (include battery flows if present) ---
labels = ["Self-consumed", "Exported", "Imported"]
values = [battery['total_self_consumed'], battery['total_exported'], battery['total_imported']]
if battery["use_battery"]:
    labels += ["Battery Charge", "Battery Discharge"]
    values += [battery['total_battery_charge'], battery['total_battery_discharge']]
plt.figure(figsize=(6, 6))
plt.pie(values, labels=labels, autopct="%1.1f%%", startangle=90)
plt.title("Annual Energy Breakdown")
plt.tight_layout()
plot3_path = os.path.join(plots_dir, "energy_breakdown.png")
plt.savefig(plot3_path)
plt.close()

# --- Calculate financial impact for each section in the pie chart ---
energy_rate = 0.315823
# Self-consumed: value is what would have been paid to the grid
self_consumed_value = battery['total_self_consumed'] * energy_rate
# Exported: earnings from feed-in tariff (weighted average)
exported_value = (df['exported'] * df['feedin_rate']).sum()
# Imported: cost
imported_value = battery['total_imported'] * energy_rate
# Battery charge/discharge: not directly financial, but show kWh
battery_charge_value = battery['total_battery_charge'] if battery['use_battery'] else 0
battery_discharge_value = battery['total_battery_discharge'] if battery['use_battery'] else 0

# --- Create Word document ---
def set_metric_col_width_and_border(table):
    # Set first column width and right border
    for row in table.rows:
        cell = row.cells[0]
        cell.width = Cm(9.45)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        right = OxmlElement('w:right')
        right.set(qn('w:val'), 'single')
        right.set(qn('w:sz'), '4')
        right.set(qn('w:space'), '0')
        right.set(qn('w:color'), 'auto')
        tcBorders.append(right)
        tcPr.append(tcBorders)

report_dir = "solar_reports"
os.makedirs(report_dir, exist_ok=True)
report_path = os.path.join(report_dir, "solar_system_report.docx")
doc = Document()
doc.add_heading("Solar System Financial & Energy Report", 0)

doc.add_heading("Summary", level=1)
for i in range(1, 6):
    para = summary_json.get(f'paragraph_{i}', '').strip()
    if para:
        doc.add_paragraph(para)
# Add a page break after the summary
if doc.paragraphs:
    doc.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)

doc.add_heading("Energy Summary Table (Solar Only)", level=2)
table1 = doc.add_table(rows=1, cols=2)
table1.style = 'Light List Accent 1'
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Value'
for k, v in energy_rows_solar:
    row_cells = table1.add_row().cells
    row_cells[0].text = k
    row_cells[1].text = v
set_metric_col_width_and_border(table1)

doc.add_heading("Cost Summary Table (Solar Only)", level=2)
table2 = doc.add_table(rows=1, cols=2)
table2.style = 'Light List Accent 2'
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Value'
for k, v in cost_rows_solar:
    row_cells = table2.add_row().cells
    row_cells[0].text = k
    row_cells[1].text = v
set_metric_col_width_and_border(table2)

doc.add_heading("Energy Summary Table (Solar + Battery)", level=2)
table3 = doc.add_table(rows=1, cols=2)
table3.style = 'Light List Accent 1'
hdr_cells = table3.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Value'
for k, v in energy_rows_battery:
    row_cells = table3.add_row().cells
    row_cells[0].text = k
    row_cells[1].text = v
set_metric_col_width_and_border(table3)

doc.add_heading("Cost Summary Table (Solar + Battery)", level=2)
table4 = doc.add_table(rows=1, cols=2)
table4.style = 'Light List Accent 2'
hdr_cells = table4.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Value'
for k, v in cost_rows_battery:
    row_cells = table4.add_row().cells
    row_cells[0].text = k
    row_cells[1].text = v
set_metric_col_width_and_border(table4)

doc.add_heading("Assumptions & Considerations", level=1)
doc.add_paragraph(assumptions_text)

# --- Landscape section for graphs ---
section = doc.add_section(WD_ORIENT.LANDSCAPE)
section.orientation = WD_ORIENT.LANDSCAPE
new_width, new_height = section.page_height, section.page_width
section.page_width = new_width
section.page_height = new_height

doc.add_heading("Graphs", level=1)
doc.add_paragraph("Hourly Energy Flows (First Week):")
pic1 = doc.add_picture(plot1_path, width=Inches(9))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)
doc.add_paragraph("Monthly Energy Totals:")
pic2 = doc.add_picture(plot2_path, width=Inches(9))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)
doc.add_paragraph("Annual Energy Breakdown:")
pic3 = doc.add_picture(plot3_path, width=Inches(6.3))  # 10% smaller
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add table for financial impact of each section in the pie chart
doc.add_paragraph("Annual Energy Breakdown: Financial Impact")
impact_table = doc.add_table(rows=1, cols=3)
impact_table.style = 'Light List Accent 1'
impact_table.cell(0, 0).text = 'Section'
impact_table.cell(0, 1).text = 'kWh'
impact_table.cell(0, 2).text = 'Financial Impact ($)'
impact_rows = [
    ("Self-consumed", f"{battery['total_self_consumed']:.2f}", f"${self_consumed_value:,.2f}"),
    ("Exported", f"{battery['total_exported']:.2f}", f"${exported_value:,.2f}"),
    ("Imported", f"{battery['total_imported']:.2f}", f"-${imported_value:,.2f}"),
]
if battery["use_battery"]:
    impact_rows.append(("Battery Charge", f"{battery_charge_value:.2f}", "-"))
    impact_rows.append(("Battery Discharge", f"{battery_discharge_value:.2f}", "-"))
for row in impact_rows:
    cells = impact_table.add_row().cells
    cells[0].text = row[0]
    cells[1].text = row[1]
    cells[2].text = row[2]
set_metric_col_width_and_border(impact_table)

doc.add_paragraph("\nReport generated automatically.")

doc.save(report_path)
print(f"\nReport generated: {report_path}") 